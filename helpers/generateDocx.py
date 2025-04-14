from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

def add_bold_heading(paragraph, text):
    run = paragraph.add_run(text.replace("#####", "").strip())
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_table(doc, table_lines):
    headers = [h.strip() for h in table_lines[0].strip().split('|') if h.strip()]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.allow_autofit = True

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for para in hdr_cells[i].paragraphs:
            run = para.runs[0]
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

    for row_line in table_lines[1:]:
        if not row_line.strip() or '---' in row_line:
            continue
        cols = [c.strip() for c in row_line.strip().split('|') if c.strip()]
        if len(cols) != len(headers):
            continue
        row_cells = table.add_row().cells
        for i, val in enumerate(cols):
            for para in row_cells[i].paragraphs:
                run = para.add_run(val)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("\n")  # Bottom padding after table

def generate_docx(markdown_text: str) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    lines = markdown_text.splitlines()
    table_buffer = []
    in_table = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("#####") or re.match(r'^\d+\.', stripped):
            if table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            para = doc.add_paragraph()
            add_bold_heading(para, stripped.replace("*", ""))

        elif "|" in stripped and "---" in stripped:
            continue  # Skip markdown table separators

        elif "|" in stripped:
            in_table = True
            table_buffer.append(stripped)

        elif in_table and not stripped:
            add_table(doc, table_buffer)
            table_buffer = []
            in_table = False

        elif line.startswith("```"):
            continue

        else:
            if table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            para = doc.add_paragraph()
            cleaned_text = stripped.replace("**", "")
            run = para.add_run(cleaned_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

    if table_buffer:
        add_table(doc, table_buffer)

    # Save as byte stream
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
