from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_bold_heading(paragraph, text):
    run = paragraph.add_run(text.replace("#####", "").strip())
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_table(doc, table_lines):
    headers = [h.strip() for h in table_lines[0].strip().split('|') if h.strip()]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.allow_autofit = True

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for para in hdr_cells[i].paragraphs:
            run = para.runs[0]
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

    for row_line in table_lines[1:]:
        if not row_line.strip() or '---' in row_line:
            continue
        cols = [c.strip() for c in row_line.strip().split('|') if c.strip()]
        if len(cols) != len(headers):
            continue
        row_cells = table.add_row().cells
        for i, val in enumerate(cols):
            for para in row_cells[i].paragraphs:
                run = para.add_run(val)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("\n")  # Bottom padding after table


def markdown_to_docx(input_file, output_file):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    with open(input_file, "r", encoding="utf-8") as f:
        lines = f.readlines()

    table_buffer = []
    in_table = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("#####") or (stripped[:2].isdigit() and stripped[2] == '.'):
            if table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            para = doc.add_paragraph()
            add_bold_heading(para, stripped)

        elif "|" in line and "---" in line:
            continue  # skip separator line

        elif "|" in line:
            in_table = True
            table_buffer.append(line)

        elif in_table and not stripped:
            add_table(doc, table_buffer)
            table_buffer = []
            in_table = False

        else:
            if table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            para = doc.add_paragraph()
            if stripped.startswith("**") and stripped.endswith("**"):
                run = para.add_run(stripped.strip("*"))
                run.bold = True
            else:
                run = para.add_run(stripped)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

    if table_buffer:
        add_table(doc, table_buffer)

    doc.save(output_file)


# Example usage
markdown_to_docx("input.txt", "AXA_Report.docx")