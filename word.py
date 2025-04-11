from docx import Document
from docx.shared import Inches

# Create a new Word Document
doc = Document()
doc.add_heading("AXA Financial Data", level=1)

# Table 1
doc.add_heading("Table 1", level=2)
table1_data = [
    ["", "AXA (CAC 40)"],
    ["Ticker", "CS FP Equity"],
    ["YTD Performance", "0.1%"],
    ["EQY_REC_CONS", "4.56"],
    ["BDVD_PROJ_DIV_AMT", "2.3"],
    ["IS_DIV_PER_SHARE", "1.98"],
    ["LOW_52WEEK", "29.04"],
    ["HIGH_52WEEK", "40.51"],
    ["Scoring", "6.82"],
]

table1 = doc.add_table(rows=1, cols=2)
table1.style = "Table Grid"
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = table1_data[0][0]
hdr_cells[1].text = table1_data[0][1]

for row_data in table1_data[1:]:
    row_cells = table1.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]

# Table 2
doc.add_heading("Table 2", level=2)
table2_data = [
    ["", "AXA (CAC 40)"],
    ["Dividende par action 2024", "2,15 EUR"],
    ["Dividende par action est 2025", "2,32 EUR"],
    ["Dividende par action est 2026", "2,50 EUR"],
    ["PER 2024", "10,47"],
    ["PER 2025", "9,82"],
    ["PER 2026", "9,09"],
    ["1year Perf", "+10,93%"],
    ["3 year perf", "+41,70%"],
    ["5 year perf", "+142,35%"],
    ["Objectif de cours", "42,49"],
]

table2 = doc.add_table(rows=1, cols=2)
table2.style = "Table Grid"
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = table2_data[0][0]
hdr_cells[1].text = table2_data[0][1]

for row_data in table2_data[1:]:
    row_cells = table2.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]

# References
doc.add_heading("References", level=2)
doc.add_paragraph("AXA:")
doc.add_paragraph("Main URL: https://www.boursorama.com/cours/1rPCS/")
doc.add_paragraph("Consensus URL: https://www.boursorama.com/cours/consensus/1rPCS/")

# Save document
doc.save("axa_financial_data.docx")
